import fitz
from docx import Document
import subprocess
import random
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import base64
from IPython.display import HTML
from pdf2image import convert_from_path
import tempfile
import viz
import ta_pipeline as ta


def toggle_example_text(checked):
    return ta.load_doc("assets/Examples/Tony Avidnote.docx", path=True) if checked else ""


def load_html_file(filepath="network.html"):
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except FileNotFoundError:
        return "<p><strong>Visualization file not found.</strong></p>"


def export_html(html_obj, highlight_type="Output"):
    if html_obj is None:
        return None

    html_str = html_obj.value if hasattr(html_obj, 'value') else html_obj

    path = f"{highlight_type}.html"

    with open(path, "w", encoding="utf-8") as f:
        f.write(html_str)
    return path


def download_theme_table(combined_dict_state):
    theme_df = viz.dict_to_table(combined_dict_state)
    csv_path = "Theme Table.csv"
    theme_df.to_csv(csv_path, index=False)
    return csv_path


def rgb_to_hex(rgb_tuple):
    return "%02x%02x%02x" % rgb_tuple

def highlight_run_bg(run, hex_color):
    """
    Apply background shading to a run (hex format like 'ffccee').
    """
    rPr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    rPr.append(shd)

def convert_docx_to_pdf(input_path, output_folder="/content"):
    subprocess.run([
        "libreoffice",
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_folder,
        input_path
    ])

def get_dynamic_colors(n, seed = 1):
    random.seed(seed)
    colors = []

    for _ in range(n):
      r = random.randint(0, 200)
      g = random.randint(0, 200)
      b = random.randint(0, 200)
      color = f'#{r:02x}{g:02x}{b:02x}'
      colors.append(color)

    return colors

def highlight_quotes_in_docx(input_path, output_path, combined_dict):
    doc = Document(input_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    theme_colors = {}
    color_palette = get_dynamic_colors(len(combined_dict))

    # Build map of quote -> theme + color
    quote_map = {}
    for idx, (theme, info) in enumerate(combined_dict.items()):
        color = color_palette[idx % len(color_palette)]
        theme_colors[theme] = color
        theme_info = info[1]
        for code, lst in theme_info.items():
            for tup in lst:
                quote_map[tup[0]] = (theme, color)

    for para in doc.paragraphs:
        for quote, (theme, rgb) in quote_map.items():
            if quote in para.text:
                new_runs = []
                parts = para.text.split(quote)
                para.clear()  # remove original run

                for i, part in enumerate(parts):
                    run = para.add_run(part)
                    new_runs.append(run)

                    if i < len(parts) - 1:
                        run = para.add_run(quote)
                        font = run.font
                        highlight_color = rgb

                        # Apply background highlight via XML
                        highlight_run_bg(run, highlight_color)
                        new_runs.append(run)

    doc.save(output_path)


def get_quote_positions(pdf_path: str, combined_dict: list[dict]):
    doc = fitz.open(pdf_path)
    boxes = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        for theme, theme_data in combined_dict.items():
            theme_info = theme_data[1]
            for code, lst in theme_info.items():
                for tup in lst:
                  quote = tup[0]
                  text_instances = page.search_for(quote)
                  for inst in text_instances:
                      boxes.append({
                          "page": page_num,
                          "theme": theme,
                          "quote": quote,
                          "bbox": inst  # (x0, y0, x1, y1)
                      })
    return boxes


def scale_image_pdf(path_to_pdf: str, path_to_image: str, combined_dict: dict) -> list[dict]:
    doc = fitz.open(path_to_pdf)
    page = doc[0]
    pdf_width, pdf_height = page.rect.width, page.rect.height

    im = Image.open(path_to_image)
    img_width, img_height = im.size

    scale_x = img_width / pdf_width
    scale_y = img_height / pdf_height

    boxes = get_quote_positions(path_to_pdf, combined_dict)
    scaled_boxes = []

    for b in boxes:
        if b["page"] != 0:
            continue
        x0, y0, x1, y1 = b["bbox"]
        scaled_box = {
            "x": x0 * scale_x,
            "y": y0 * scale_y,
            "w": (x1 - x0) * scale_x,
            "h": (y1 - y0) * scale_y,
            "theme": b["theme"],
            "quote": b["quote"]
        }
        scaled_boxes.append(scaled_box)

    return scaled_boxes


def show_interactive_overlay(image_path: str, scaled_boxes: list[dict], combined_dict: list[dict], seed: int=42):
    dynamic_color_palette = get_dynamic_colors(len(combined_dict))
    theme_colors = {theme: dynamic_color_palette[i] for i,theme in enumerate(combined_dict.keys())}

    with open(image_path, "rb") as f:
        img_data = base64.b64encode(f.read()).decode("utf-8")

    divs = ""
    for b in scaled_boxes:
        color = "rgb" + str(theme_colors[b["theme"]])
        divs += f"""
        <div class='tooltip' style='
            position: absolute;
            left: {b["x"]}px;
            top: {b["y"]}px;
            width: {b["w"]}px;
            height: {b["h"]}px;
            background-color: {color};
            opacity: 0.3;
            border-radius: 4px;
        '>
            <span class='tooltiptext'>{b["theme"]}</span>
        </div>
        """

    html = f"""
    <style>
    .tooltip .tooltiptext {{
        visibility: hidden;
        width: auto;
        background-color: rgba(0, 0, 0, 0.9);
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 3px 6px;
        position: absolute;
        z-index: 1;
        top: -5px;
        left: 105%;
        font-size: 12px;
    }}
    .tooltip:hover .tooltiptext {{
        visibility: visible;
    }}
    </style>
    <div style="position: relative; display: inline-block;">
        <img src="data:image/png;base64,{img_data}" />
        {divs}
    </div>
    """
    display(HTML(html))